from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse
import re

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from src.discovery.fifa_block_data import FIFABlockDataExtractor, ExtractedBlockData
from src.reports.variable_dictionary import build_variable_dictionary


def _safe_name(v: str) -> str:
    return "_".join(v.replace("/", " ").split())


def json_safe(value):
    import json
    try:
        import numpy as np
        if isinstance(value, np.ndarray): value = value.tolist()
        elif isinstance(value, np.generic): value = value.item()
    except ImportError:
        pass
    if isinstance(value, (dict, list, tuple, set)):
        if isinstance(value, set): value = sorted(value, key=str)
        return json.dumps(value, ensure_ascii=False, default=str)
    return value


def _clean_player_name(value):
    if not isinstance(value, str): return value
    # FIFA concatena a veces nombre + código de país duplicado + posición.
    return re.sub(r"([A-Z]{3})\1(GK|DF|DEF|MF|MID|FW|FWD)$", "", value).strip()


def _prepare_for_excel(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [str(c).strip() for c in out.columns]
    technical = re.compile(r"(^|\.)(id|identifier|uuid|slug|pageurl|resource|image|asset|theme|document|audio)(\.|$)", re.I)
    keep = [c for c in out.columns if not technical.search(str(c))]
    if keep:
        out = out[keep]
    for col in out.columns:
        if str(col).lower() in {"player", "name", "playername", "shortname"}:
            out[col] = out[col].map(_clean_player_name)
        out[col] = out[col].map(json_safe)
    return out


def _is_missing_scalar(value) -> bool:
    if isinstance(value, (dict, list, tuple, set)): return False
    try:
        result = pd.isna(value)
        return bool(result) if not hasattr(result, "size") else bool(result) if result.size == 1 else False
    except (TypeError, ValueError):
        return False


def generate_excel(block: dict, data: ExtractedBlockData, folder: Path) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    out = folder / f"{_safe_name(block['name'])}.xlsx"
    summary = pd.DataFrame([
        {"Bloque": block["name"], "Fuente": block["url"], "Dominio": urlparse(block["url"]).netloc,
         "Fecha": datetime.now().strftime("%d/%m/%Y %H:%M"), "Modo": data.discovery_mode,
         "Tablas útiles": len(data.tables), "Registros útiles": data.row_count, "Error": data.error or ""}
    ])
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        summary.to_excel(writer, sheet_name="Resumen", index=False)
        for i, df in enumerate(data.tables, 1):
            name = "Datos" if len(data.tables) == 1 else f"Datos_{i}"
            _prepare_for_excel(df).to_excel(writer, sheet_name=name[:31], index=False)
        metadata = pd.DataFrame([
            {"campo": "Página analizada", "valor": block["url"]},
            {"campo": "Bloque", "valor": block["name"]},
            {"campo": "Descripción", "valor": block["description"]},
            {"campo": "Modo de extracción", "valor": data.discovery_mode},
            {"campo": "Criterio", "valor": "Solo tablas deportivas relevantes; contenido web, menús, enlaces y endpoints excluidos."},
        ])
        metadata.to_excel(writer, sheet_name="Metadatos", index=False)
        all_columns = []
        for df in data.tables:
            for col in _prepare_for_excel(df).columns:
                if str(col) not in all_columns:
                    all_columns.append(str(col))
        build_variable_dictionary(all_columns).to_excel(writer, sheet_name="Diccionario_variables", index=False)
        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            ws.sheet_view.showGridLines = False
            for cell in ws[1]:
                cell.font = cell.font.copy(bold=True)
            for col in ws.columns:
                width = min(max(len(str(json_safe(cell.value))) if cell.value is not None else 0 for cell in col) + 2, 42)
                ws.column_dimensions[col[0].column_letter].width = max(width, 10)
    return out


def _add_dataframe(doc: Document, df: pd.DataFrame, title: str, limit: int = 150) -> None:
    doc.add_heading(title, level=1)
    if df.empty:
        doc.add_paragraph("No se recuperaron registros útiles.")
        return
    display = _prepare_for_excel(df).head(limit).iloc[:, :12]
    table = doc.add_table(rows=1, cols=len(display.columns)); table.style = "Table Grid"
    for i, col in enumerate(display.columns): table.rows[0].cells[i].text = str(col)
    for _, row in display.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if _is_missing_scalar(value) else str(json_safe(value))[:300]
    if len(df) > limit:
        doc.add_paragraph(f"Se muestran {limit} de {len(df)} registros. El Excel contiene la tabla completa.")


def generate_word(block: dict, data: ExtractedBlockData, folder: Path) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    out = folder / f"{_safe_name(block['name'])}.docx"
    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE; sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = sec.bottom_margin = Cm(1.4); sec.left_margin = sec.right_margin = Cm(1.4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SportsLab-Analytics"); r.bold = True; r.font.size = Pt(20)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(block["name"]); r.bold = True; r.font.size = Pt(16)
    doc.add_paragraph(block["description"])
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Indicador"; table.rows[0].cells[1].text = "Resultado"
    for key, value in [("Fuente", block["url"]), ("Fecha", datetime.now().strftime("%d/%m/%Y %H:%M")),
                       ("Modo", data.discovery_mode), ("Tablas útiles", len(data.tables)),
                       ("Registros útiles", data.row_count), ("Error", data.error or "Ninguno")]:
        cells = table.add_row().cells; cells[0].text = str(key); cells[1].text = str(value)
    for i, df in enumerate(data.tables, 1): _add_dataframe(doc, df, "Datos" if len(data.tables)==1 else f"Datos {i}")
    if not data.tables:
        doc.add_heading("Resultado", level=1); doc.add_paragraph("No se identificaron tablas deportivas válidas.")
    doc.save(out); return out


def generate_reports(block: dict, base_results: Path) -> dict:
    folder = base_results / _safe_name(block["name"])
    data = FIFABlockDataExtractor().extract(block)
    excel = generate_excel(block, data, folder / "Excel")
    word = generate_word(block, data, folder / "Word")
    return {"block": block["name"], "excel": str(excel), "word": str(word),
            "tables": len(data.tables), "rows": data.row_count, "mode": data.discovery_mode, "error": data.error}
